#!/usr/bin/env python3
"""
COBS Bread Bakery Deep Research Application
Uses Google Deep Research API to analyze bakery reviews across all social media platforms.
"""

import time
import os
import sys
import argparse
from datetime import datetime
from typing import Optional
from pathlib import Path

from dotenv import load_dotenv
from google import genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class COBSBreadResearcher:
    """Deep research engine for COBS Bread bakery reviews."""

    AGENT_ID = "deep-research-pro-preview-12-2025"
    MAX_POLL_TIME = 3600  # 60 minutes max
    POLL_INTERVAL = 15  # seconds

    def __init__(self, api_key: Optional[str] = None):
        """Initialize the researcher with Google API credentials."""
        if api_key:
            os.environ["GOOGLE_API_KEY"] = api_key

        if not os.environ.get("GOOGLE_API_KEY"):
            raise ValueError(
                "Google API key required. Set GOOGLE_API_KEY environment variable "
                "or pass api_key parameter."
            )

        self.client = genai.Client()
        self.results = {}

    def _build_research_prompt(self, bakery_location: str) -> str:
        """Build comprehensive research prompt for COBS Bread bakery analysis."""
        return f"""
You are conducting an exhaustive deep research analysis of customer reviews for the COBS Bread bakery located at: {bakery_location}

Your task is to find and analyze ALL available reviews across EVERY social media platform and review site. Be extremely thorough and comprehensive.

## PLATFORMS TO SEARCH (search ALL of these):
- Google Reviews / Google Maps
- Yelp
- Facebook (page reviews and comments)
- Instagram (mentions, tagged posts, comments)
- Twitter/X (mentions, hashtags)
- TikTok (mentions, reviews, videos)
- Reddit (r/vancouver, r/calgary, r/toronto, local subreddits, food subreddits)
- TripAdvisor
- Zomato
- DoorDash reviews
- UberEats reviews
- Skip The Dishes reviews
- Local food blogs and review sites
- News articles mentioning this location
- YouTube reviews and mentions
- LinkedIn (any business mentions)
- NextDoor app mentions
- Local community forums
- Food critic reviews
- Franchise review sites

## ANALYSIS FRAMEWORK:

### SECTION 1: 5-STAR REVIEWS (Excellent - Detailed Analysis)
For every 5-star review found, extract and analyze:
1. **Beloved Products**: List EVERY product mentioned positively with specific details
   - Bread types (sourdough, rye, whole wheat, etc.)
   - Pastries (croissants, danishes, muffins, etc.)
   - Sandwiches and prepared foods
   - Seasonal/special items
   - Frequency of mention for each product

2. **Customer Experience Highlights**:
   - Staff interactions and service quality
   - Store atmosphere and cleanliness
   - Wait times and efficiency
   - Product freshness observations
   - Packaging quality
   - In-store experience

3. **Discovery Journey**:
   - How customers found this bakery (word of mouth, walking by, online search, etc.)
   - What made them first try it
   - How long they've been customers
   - Referral patterns

4. **Loyalty Indicators**:
   - Repeat customer patterns
   - What keeps them coming back
   - Comparisons to other bakeries
   - Brand affinity statements

5. **Value Perception**:
   - Price satisfaction
   - Quality-to-price ratio comments
   - Deals and promotions mentioned

### SECTION 2: 4-STAR REVIEWS (Good but with reservations - Balanced Analysis)
Analyze from DUAL PERSPECTIVES:

**A) Customer/Reviewer Perspective**:
1. **What They Loved**:
   - Specific products praised
   - Positive experiences
   - Strengths identified

2. **What Held Back 5 Stars**:
   - Specific criticisms
   - Products that disappointed
   - Service issues
   - Suggestions for improvement

3. **Conditional Recommendations**:
   - "Great for X but not for Y" statements
   - Situational preferences

**B) COBS Bread Franchisor Perspective**:
1. **Operational Insights**:
   - Consistency issues flagged
   - Training opportunities identified
   - Product quality variations
   - Peak time challenges

2. **Competitive Intelligence**:
   - Comparisons to competitors
   - Market positioning feedback
   - Pricing perception

3. **Growth Opportunities**:
   - Unmet customer needs
   - Product suggestions
   - Service improvements needed

### SECTION 3: 3-STAR AND BELOW (Critical Analysis)
For every review rated 3 stars or lower, deeply analyze:

1. **Problematic Products** (CRITICAL - Be Exhaustive):
   - List EVERY product mentioned negatively
   - Specific issues (staleness, taste, texture, appearance)
   - Frequency of complaints per product
   - Pattern analysis across reviews

2. **Service Failures**:
   - Staff behavior issues
   - Wait time complaints
   - Order accuracy problems
   - Customer service recovery (or lack thereof)

3. **Environmental Issues**:
   - Cleanliness concerns
   - Store layout problems
   - Parking/accessibility issues
   - Noise/crowding complaints

4. **Value Disappointments**:
   - Price complaints
   - Portion size issues
   - Quality-to-price mismatch

5. **Discovery Process Analysis**:
   - How these dissatisfied customers found the bakery
   - What expectations were set vs. reality
   - Impact on word-of-mouth

6. **Severity Assessment**:
   - One-time issues vs. recurring problems
   - Seasonal patterns in complaints
   - Response from bakery (if any)

### SECTION 4: COMPETITIVE LANDSCAPE
- How does this COBS location compare to:
  - Other COBS Bread locations nearby
  - Local independent bakeries
  - Chain competitors (Tim Hortons, Panera, etc.)
  - Grocery store bakeries

### SECTION 5: TREND ANALYSIS
- Review sentiment over time
- Seasonal patterns
- Post-COVID changes
- New product reception
- Staff/management change indicators

### SECTION 6: SOCIAL MEDIA SENTIMENT
- Hashtag analysis
- Viral moments (positive or negative)
- Influencer mentions
- User-generated content themes

### SECTION 7: ACTIONABLE INSIGHTS SUMMARY
Provide specific, actionable recommendations for:
1. Products to highlight/promote
2. Products to improve or consider removing
3. Service training priorities
4. Marketing opportunities
5. Competitive positioning strategies

## OUTPUT FORMAT:
- Be extremely detailed and thorough
- Include specific quotes from reviews where possible
- Provide counts/statistics where available
- Organize clearly by section
- Include the source platform for each insight
- Note the approximate date range of reviews analyzed
- Flag any data limitations or gaps in available information

Remember: This research will inform critical business decisions. Leave no stone unturned.
"""

    def conduct_research(self, bakery_location: str, verbose: bool = True) -> str:
        """
        Conduct deep research on the specified COBS Bread bakery.

        Args:
            bakery_location: Full address or location description of the bakery
            verbose: Whether to print progress updates

        Returns:
            Complete research report as a string
        """
        prompt = self._build_research_prompt(bakery_location)

        if verbose:
            print(f"\n{'='*60}")
            print("COBS Bread Deep Research - Initiating")
            print(f"{'='*60}")
            print(f"Target Location: {bakery_location}")
            print(f"Agent: {self.AGENT_ID}")
            print(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            print(f"{'='*60}\n")
            print("Launching deep research agent...")
            print("This may take 10-30 minutes for comprehensive analysis.\n")

        # Create the research interaction
        interaction = self.client.interactions.create(
            input=prompt,
            agent=self.AGENT_ID,
            background=True
        )

        interaction_id = interaction.id
        if verbose:
            print(f"Research task created. ID: {interaction_id}")
            print("Polling for results...\n")

        # Poll for results
        start_time = time.time()
        last_status = None

        while True:
            elapsed = time.time() - start_time

            if elapsed > self.MAX_POLL_TIME:
                raise TimeoutError(
                    f"Research exceeded maximum time of {self.MAX_POLL_TIME/60} minutes"
                )

            interaction = self.client.interactions.get(interaction_id)

            if interaction.status != last_status:
                last_status = interaction.status
                if verbose:
                    mins = int(elapsed // 60)
                    secs = int(elapsed % 60)
                    print(f"[{mins:02d}:{secs:02d}] Status: {interaction.status}")

            if interaction.status == "completed":
                if verbose:
                    print(f"\nResearch completed in {elapsed/60:.1f} minutes!")

                # Extract the final report
                if interaction.outputs:
                    report = interaction.outputs[-1].text
                    self.results[bakery_location] = report
                    return report
                else:
                    raise ValueError("Research completed but no output received")

            elif interaction.status == "failed":
                error_msg = getattr(interaction, 'error', 'Unknown error')
                raise RuntimeError(f"Research failed: {error_msg}")

            time.sleep(self.POLL_INTERVAL)

    def generate_word_document(
        self,
        bakery_location: str,
        report_content: str,
        output_path: Optional[str] = None
    ) -> str:
        """
        Generate a formatted Word document from the research report.

        Args:
            bakery_location: The bakery location researched
            report_content: The research report content
            output_path: Optional custom output path

        Returns:
            Path to the generated document
        """
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_location = "".join(c if c.isalnum() or c in ' -_' else '_' for c in bakery_location)[:50]
            output_path = f"COBS_Research_{safe_location}_{timestamp}.docx"

        doc = Document()

        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Title
        title = doc.add_heading('COBS Bread Bakery', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading('Comprehensive Review Analysis Report', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata section
        doc.add_paragraph()
        meta = doc.add_paragraph()
        meta.add_run('Location: ').bold = True
        meta.add_run(bakery_location)

        meta2 = doc.add_paragraph()
        meta2.add_run('Generated: ').bold = True
        meta2.add_run(datetime.now().strftime('%B %d, %Y at %I:%M %p'))

        meta3 = doc.add_paragraph()
        meta3.add_run('Research Engine: ').bold = True
        meta3.add_run('Google Deep Research API (Gemini)')

        doc.add_paragraph()
        doc.add_paragraph('_' * 70)
        doc.add_paragraph()

        # Process and add the report content
        self._add_formatted_content(doc, report_content)

        # Footer
        doc.add_paragraph()
        doc.add_paragraph('_' * 70)
        footer = doc.add_paragraph()
        footer.add_run('Disclaimer: ').bold = True
        footer.add_run(
            'This report is generated from publicly available reviews and social media content. '
            'All insights should be verified independently before making business decisions.'
        )

        # Save the document
        doc.save(output_path)
        return output_path

    def _add_formatted_content(self, doc: Document, content: str):
        """Add formatted content to the Word document, parsing markdown-like formatting."""
        lines = content.split('\n')
        current_list_level = 0

        for line in lines:
            stripped = line.strip()

            if not stripped:
                doc.add_paragraph()
                continue

            # Handle headers
            if stripped.startswith('######'):
                doc.add_heading(stripped[6:].strip(), level=6)
            elif stripped.startswith('#####'):
                doc.add_heading(stripped[5:].strip(), level=5)
            elif stripped.startswith('####'):
                doc.add_heading(stripped[4:].strip(), level=4)
            elif stripped.startswith('###'):
                doc.add_heading(stripped[3:].strip(), level=3)
            elif stripped.startswith('##'):
                doc.add_heading(stripped[2:].strip(), level=2)
            elif stripped.startswith('#'):
                doc.add_heading(stripped[1:].strip(), level=1)

            # Handle bullet points
            elif stripped.startswith('- ') or stripped.startswith('* '):
                para = doc.add_paragraph(stripped[2:], style='List Bullet')

            # Handle numbered lists
            elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
                para = doc.add_paragraph(stripped[2:].strip(), style='List Number')

            # Handle bold text markers
            elif stripped.startswith('**') and stripped.endswith('**'):
                para = doc.add_paragraph()
                para.add_run(stripped[2:-2]).bold = True

            # Regular paragraph
            else:
                para = doc.add_paragraph()
                # Handle inline bold
                self._add_inline_formatting(para, stripped)

    def _add_inline_formatting(self, paragraph, text: str):
        """Handle inline bold and italic formatting."""
        import re

        # Pattern for **bold** text
        pattern = r'\*\*([^*]+)\*\*'
        parts = re.split(pattern, text)

        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Regular text
                paragraph.add_run(part)
            else:
                # Bold text
                paragraph.add_run(part).bold = True


def main():
    """Main entry point for the COBS Bread research application."""
    # Load environment variables from .env file if it exists
    env_path = Path(__file__).parent / '.env'
    if env_path.exists():
        load_dotenv(env_path)

    parser = argparse.ArgumentParser(
        description='Deep Research Tool for COBS Bread Bakery Reviews',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s "COBS Bread, 123 Main Street, Vancouver, BC"
  %(prog)s "COBS Bread Bakery, Calgary, AB" --output my_report.docx
  %(prog)s "COBS Bread, Toronto, ON" --api-key YOUR_API_KEY

Note: Set GOOGLE_API_KEY environment variable or use --api-key flag.
        """
    )

    parser.add_argument(
        'location',
        nargs='?',
        help='COBS Bread bakery location (address or description)'
    )

    parser.add_argument(
        '--output', '-o',
        help='Output path for the Word document'
    )

    parser.add_argument(
        '--api-key', '-k',
        help='Google API key (or set GOOGLE_API_KEY env variable)'
    )

    parser.add_argument(
        '--quiet', '-q',
        action='store_true',
        help='Suppress progress output'
    )

    args = parser.parse_args()

    # Interactive mode if no location provided
    if not args.location:
        print("\n" + "="*60)
        print("  COBS Bread Bakery - Deep Research Tool")
        print("="*60)
        print("\nThis tool conducts exhaustive research on COBS Bread bakery")
        print("reviews across all social media platforms and review sites.\n")

        location = input("Enter the COBS Bread bakery location\n(e.g., 'COBS Bread, 123 Main St, Vancouver, BC'): ").strip()

        if not location:
            print("Error: Location is required.")
            sys.exit(1)
    else:
        location = args.location

    try:
        # Initialize researcher
        researcher = COBSBreadResearcher(api_key=args.api_key)

        # Conduct research
        report = researcher.conduct_research(
            bakery_location=location,
            verbose=not args.quiet
        )

        # Generate Word document
        doc_path = researcher.generate_word_document(
            bakery_location=location,
            report_content=report,
            output_path=args.output
        )

        print(f"\n{'='*60}")
        print("RESEARCH COMPLETE")
        print(f"{'='*60}")
        print(f"\nWord document saved to: {doc_path}")
        print(f"Report length: {len(report):,} characters")
        print("\nYou can now open the document in Microsoft Word or any")
        print("compatible application to review the findings.\n")

    except ValueError as e:
        print(f"\nConfiguration Error: {e}")
        print("Please set your Google API key:")
        print("  export GOOGLE_API_KEY='your-api-key'")
        print("  or use: --api-key YOUR_KEY")
        sys.exit(1)

    except TimeoutError as e:
        print(f"\nTimeout Error: {e}")
        sys.exit(1)

    except Exception as e:
        print(f"\nError: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
