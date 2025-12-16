#!/usr/bin/env python3
"""
COBS Bread Research - Flask Web Application
Provides a web interface for the deep research tool.
"""

import os
import time
import uuid
import json
import threading
import requests
from datetime import datetime
from pathlib import Path
from typing import Optional

from flask import Flask, render_template, request, jsonify, send_file, Response
from dotenv import load_dotenv
from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')

# File-based task storage for persistence across restarts
TASKS_FILE = Path('tasks.json')
OUTPUTS_DIR = Path('outputs')

# Google Deep Research agent ID
AGENT_ID = "deep-research-pro-preview-12-2025"

# Model for Google Search grounding (Gemini 2.x)
GROUNDING_MODEL = "gemini-2.5-flash"


# =============================================================================
# PREFETCH FUNCTIONS - Get verified data before Deep Research
# =============================================================================

def find_place_id(query: str) -> str:
    """Find Google Place ID using Places API Text Search."""
    api_key = os.environ.get('GOOGLE_PLACES_API_KEY')
    if not api_key:
        return None

    url = "https://places.googleapis.com/v1/places:searchText"
    headers = {
        "Content-Type": "application/json",
        "X-Goog-Api-Key": api_key,
        "X-Goog-FieldMask": "places.id,places.displayName,places.formattedAddress"
    }
    payload = {"textQuery": query}

    try:
        response = requests.post(url, json=payload, headers=headers, timeout=30)
        if response.status_code == 200:
            data = response.json()
            places = data.get("places", [])
            if places:
                return places[0].get('id')
    except Exception as e:
        print(f"Error finding place ID: {e}")

    return None


def fetch_google_reviews(location: str) -> dict:
    """
    Fetch the 5 most recent Google reviews using Places API (Legacy).
    Legacy API supports reviews_sort=newest parameter.
    """
    api_key = os.environ.get('GOOGLE_PLACES_API_KEY')
    if not api_key:
        return {'success': False, 'error': 'No Places API key'}

    # Step 1: Find Place ID
    place_id = find_place_id(f"COBS Bread {location}")
    if not place_id:
        return {'success': False, 'error': 'Could not find place'}

    # Step 2: Fetch reviews using Legacy API (supports newest sorting)
    url = "https://maps.googleapis.com/maps/api/place/details/json"
    params = {
        "place_id": place_id,
        "fields": "name,formatted_address,formatted_phone_number,rating,user_ratings_total,opening_hours,reviews",
        "reviews_sort": "newest",
        "key": api_key
    }

    try:
        response = requests.get(url, params=params, timeout=30)
        if response.status_code != 200:
            return {'success': False, 'error': f'API error: {response.status_code}'}

        data = response.json()
        if data.get("status") != "OK":
            return {'success': False, 'error': data.get('status')}

        result = data.get("result", {})
        reviews = result.get("reviews", [])

        # Format reviews for the prompt
        formatted_reviews = []
        for r in reviews:
            formatted_reviews.append({
                'author': r.get('author_name', 'Anonymous'),
                'rating': r.get('rating', 'N/A'),
                'time': r.get('relative_time_description', 'Unknown'),
                'text': r.get('text', '')
            })

        return {
            'success': True,
            'business_name': result.get('name', 'COBS Bread'),
            'address': result.get('formatted_address', ''),
            'phone': result.get('formatted_phone_number', ''),
            'rating': result.get('rating', 'N/A'),
            'total_reviews': result.get('user_ratings_total', 0),
            'reviews': formatted_reviews,
            'place_id': place_id
        }

    except Exception as e:
        return {'success': False, 'error': str(e)}


def fetch_search_grounding_insights(location: str) -> dict:
    """
    Fetch review insights using Gemini Google Search grounding.
    Searches across Yelp, TripAdvisor, Reddit, UberEats, etc.
    """
    api_key = os.environ.get('GOOGLE_API_KEY')
    if not api_key:
        return {'success': False, 'error': 'No Gemini API key'}

    prompt = f"""
Search the web for reviews and feedback SPECIFICALLY about the COBS Bread bakery location in {location}.

IMPORTANT: Only include reviews and discussions that specifically mention this {location} location.
Do NOT include general COBS Bread reviews or reviews from other COBS Bread locations.

Search these platforms for this specific location:
- Yelp reviews for "COBS Bread {location}"
- TripAdvisor reviews for "COBS Bread {location}"
- UberEats/DoorDash reviews for "COBS Bread {location}"
- Reddit discussions mentioning "COBS Bread" AND "{location}"
- RedFlagDeals discussions mentioning "COBS Bread" AND "{location}"
- Local food blogs reviewing this specific {location} bakery

For each review or mention found, provide:
1. The platform/source name
2. The specific rating or sentiment
3. What the customer said (quote if possible)
4. Date if available

If you cannot find location-specific reviews for a platform, state "No {location}-specific reviews found on [platform]" rather than providing general COBS Bread feedback.

Focus on quality over quantity - only verified {location}-specific feedback.
"""

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GROUNDING_MODEL}:generateContent?key={api_key}"
    payload = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "tools": [{"googleSearch": {}}]
    }

    try:
        response = requests.post(url, json=payload, headers={"Content-Type": "application/json"}, timeout=90)
        if response.status_code != 200:
            return {'success': False, 'error': f'API error: {response.status_code}'}

        data = response.json()
        text = data.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")

        # Extract sources
        grounding = data.get("candidates", [{}])[0].get("groundingMetadata", {})
        sources = []
        for chunk in grounding.get("groundingChunks", []):
            web_data = chunk.get("web", {})
            if web_data:
                sources.append(web_data.get('title', ''))

        return {
            'success': True,
            'insights': text,
            'sources': sources
        }

    except Exception as e:
        return {'success': False, 'error': str(e)}


def load_tasks():
    """Load tasks from file."""
    if TASKS_FILE.exists():
        try:
            with open(TASKS_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}


def save_tasks(tasks):
    """Save tasks to file."""
    try:
        with open(TASKS_FILE, 'w') as f:
            json.dump(tasks, f, indent=2)
    except Exception as e:
        print(f"Error saving tasks: {e}")


def get_task(task_id):
    """Get a specific task."""
    tasks = load_tasks()
    return tasks.get(task_id)


def update_task(task_id, updates):
    """Update a specific task."""
    tasks = load_tasks()
    if task_id in tasks:
        tasks[task_id].update(updates)
        save_tasks(tasks)
    return tasks.get(task_id)


def create_task(task_id, task_data):
    """Create a new task."""
    tasks = load_tasks()
    tasks[task_id] = task_data
    save_tasks(tasks)
    return task_data


def build_research_prompt(bakery_location: str, google_reviews: dict = None, search_insights: dict = None) -> str:
    """Build comprehensive research prompt for COBS Bread bakery analysis."""
    today = datetime.now().strftime('%B %d, %Y')

    # Build prefetched data section
    prefetch_section = ""

    # Add Google Reviews if available
    if google_reviews and google_reviews.get('success'):
        reviews_text = ""
        for r in google_reviews.get('reviews', []):
            reviews_text += f"\n- **{r['author']}** ({r['time']}) - {r['rating']}/5 stars:\n  \"{r['text']}\"\n"

        prefetch_section += f"""
## VERIFIED GOOGLE REVIEWS DATA (from Google Places API)
**Business:** {google_reviews.get('business_name', 'COBS Bread')}
**Address:** {google_reviews.get('address', 'N/A')}
**Phone:** {google_reviews.get('phone', 'N/A')}
**Google Rating:** {google_reviews.get('rating', 'N/A')}/5 ({google_reviews.get('total_reviews', 0)} total reviews)

### 5 Most Recent Google Reviews (sorted by date):
{reviews_text}
---
"""

    # Add Search Grounding insights if available
    if search_insights and search_insights.get('success'):
        prefetch_section += f"""
## MULTI-PLATFORM REVIEW INSIGHTS (from Google Search Grounding)
The following insights were gathered from web search across multiple platforms:

{search_insights.get('insights', '')}

**Sources searched:** {', '.join(search_insights.get('sources', [])[:10])}
---
"""

    # Add instruction if prefetch data exists
    if prefetch_section:
        prefetch_section = f"""
# PRE-FETCHED VERIFIED DATA
The following data has been verified and grounded from official APIs and web search.
Use this as your foundation and expand upon it with additional deep research.

{prefetch_section}

# NOW CONDUCT ADDITIONAL DEEP RESEARCH
Expand on the above verified data by searching for more reviews and insights.
---

"""

    return f"""
You are conducting an exhaustive deep research analysis of customer reviews for the COBS Bread bakery located at: {bakery_location}
{prefetch_section}
**TODAY'S DATE: {today}**

## CRITICAL REQUIREMENTS - READ CAREFULLY:

**ONLY ACTUAL REVIEWS ALLOWED:**
1. ONLY include reviews, quotes, and data that you can directly source and verify
2. NEVER infer, estimate, guess, or imply reviews or sentiments
3. NEVER use language like "inferred," "estimated," "implied," "likely," or "typical"
4. If a section has no actual reviews found, write "No verified reviews found for this category"
5. All statistics MUST be exact counts from actual found reviews - NO estimates or projections
6. Every insight MUST have a direct quote or specific reference from an actual review
7. Do NOT include competitive analysis unless a competitor is specifically mentioned in an actual review
8. Do NOT extrapolate or generalize from reviews at other COBS locations

**QUOTE REQUIREMENTS:**
- Every claim must be backed by a direct quote from an actual review
- Include the reviewer's name/username, platform, date, and rating for each quote
- If you cannot find the exact quote, do not include the insight

**DATA INTEGRITY:**
- Rating distributions must only count actual reviews found - NO estimates
- If a platform has 0 reviews, show 0 - do not estimate or project
- Do not "fill in" sections with general brand knowledge if no specific reviews exist

Your task is to find and analyze ALL available reviews across EVERY social media platform and review site. Be extremely thorough and comprehensive. Search for the most recent reviews up to today's date.

## PLATFORMS TO SEARCH (search ALL of these):
- Google Reviews / Google Maps
- Yelp
- Facebook (page reviews and comments)
- Instagram (mentions, tagged posts, comments)
- Twitter/X (mentions, hashtags)
- TikTok (mentions, reviews, videos)
- Reddit (r/vancouver, r/calgary, r/toronto, local subreddits, food subreddits)
- TripAdvisor
- Zomato
- DoorDash reviews
- UberEats reviews
- Skip The Dishes reviews
- Local food blogs and review sites
- News articles mentioning this location
- YouTube reviews and mentions
- LinkedIn (any business mentions)
- NextDoor app mentions
- Local community forums
- Food critic reviews
- Franchise review sites

## ANALYSIS FRAMEWORK:

### SECTION 1: 5-STAR REVIEWS (Perfect Score - Detailed Analysis)
For every 5-star review found, extract and analyze:
1. **Beloved Products**: List EVERY product mentioned positively with specific details
   - Bread types (sourdough, rye, whole wheat, etc.)
   - Pastries (croissants, danishes, muffins, etc.)
   - Sandwiches and prepared foods
   - Seasonal/special items
   - Frequency of mention for each product

2. **Customer Experience Highlights**:
   - Staff interactions and service quality
   - Store atmosphere and cleanliness
   - Wait times and efficiency
   - Product freshness observations
   - Packaging quality
   - In-store experience

3. **Discovery Journey**:
   - How customers found this bakery (word of mouth, walking by, online search, etc.)
   - What made them first try it
   - How long they've been customers
   - Referral patterns

4. **Loyalty Indicators**:
   - Repeat customer patterns
   - What keeps them coming back
   - Comparisons to other bakeries
   - Brand affinity statements

5. **Value Perception**:
   - Price satisfaction
   - Quality-to-price ratio comments
   - Deals and promotions mentioned

### SECTION 2: 4-STAR REVIEWS (Excellent but not perfect - Actual Reviews Only)
For reviews rated exactly 4 stars, analyze ONLY actual quotes found:

**IMPORTANT: If no 4-star reviews are found, state "No 4-star reviews found for this location" and skip this section.**

For each 4-star review found, include:
1. **Reviewer name, platform, date**
2. **Direct quote from the review**
3. **What they praised** (from the actual review text)
4. **What they criticized** (from the actual review text)
5. **Any specific products mentioned** (positive or negative)

Do NOT infer or estimate what "might have" held back 5 stars. Only include explicit statements from actual reviews.

### SECTION 3: 3-4 STAR REVIEWS (Mixed Experience - Improvement Opportunities)
For reviews rated between 3 and 4 stars, analyze:

1. **Balanced Feedback**:
   - What worked vs. what didn't
   - Products that received mixed reviews
   - Service inconsistencies noted

2. **Specific Improvement Areas**:
   - Recurring themes in "middle ground" reviews
   - Common "almost great" experiences
   - Suggestions customers provided

3. **Conversion Opportunities**:
   - What would turn these into 5-star experiences
   - Quick wins identified
   - Systemic issues to address

### SECTION 4: 3 STARS AND BELOW (Critical Analysis)
For every review rated 3 stars or lower, deeply analyze:

1. **Problematic Products** (CRITICAL - Be Exhaustive):
   - List EVERY product mentioned negatively
   - Specific issues (staleness, taste, texture, appearance)
   - Frequency of complaints per product
   - Pattern analysis across reviews

2. **Service Failures**:
   - Staff behavior issues
   - Wait time complaints
   - Order accuracy problems
   - Customer service recovery (or lack thereof)

3. **Environmental Issues**:
   - Cleanliness concerns
   - Store layout problems
   - Parking/accessibility issues
   - Noise/crowding complaints

4. **Value Disappointments**:
   - Price complaints
   - Portion size issues
   - Quality-to-price mismatch

5. **Discovery Process Analysis**:
   - How these dissatisfied customers found the bakery
   - What expectations were set vs. reality
   - Impact on word-of-mouth

6. **Severity Assessment**:
   - One-time issues vs. recurring problems
   - Seasonal patterns in complaints
   - Response from bakery (if any)

### SECTION 5: COMPETITIVE MENTIONS (From Actual Reviews Only)
**IMPORTANT: Only include competitor mentions that appear in actual reviews for THIS location.**

If reviewers explicitly compare this COBS to other bakeries, include:
- The exact quote mentioning the competitor
- Reviewer name, platform, date
- Whether the comparison was favorable or unfavorable

**If no reviews mention competitors, write: "No competitor comparisons found in reviews for this location."**

Do NOT add general competitive analysis or market positioning unless it comes from actual customer reviews.

### SECTION 6: TREND ANALYSIS
- Review sentiment over time
- Seasonal patterns
- Post-COVID changes
- New product reception
- Staff/management change indicators

### SECTION 7: SOCIAL MEDIA MENTIONS (Verified Posts Only)
**IMPORTANT: Only include actual posts/comments you can directly reference.**

For each social media mention found, include:
- Platform (Instagram, TikTok, Twitter, Reddit, etc.)
- Username/handle
- Date of post
- Direct quote or description of content
- Link if available

**If no social media mentions are found for this specific location, write: "No verified social media mentions found for this location."**

Do NOT infer sentiment from "general engagement metrics" or assume what content might exist.

### SECTION 8: ACTIONABLE INSIGHTS SUMMARY
Provide specific, actionable recommendations for:
1. Products to highlight/promote
2. Products to improve or consider removing
3. Service training priorities
4. Marketing opportunities
5. Competitive positioning strategies

## SECTION 9: REVIEW STATISTICS & DATA SOURCES (MANDATORY - ACTUAL DATA ONLY)
**THIS SECTION IS REQUIRED - ONLY include counts from reviews you actually found and can cite.**

### 9.1 Total Reviews Analyzed
- **Total number of reviews analyzed**: [exact count of reviews you found and quoted]
- **Date range of reviews**: [earliest review date] to [latest review date]

**CRITICAL: Only count reviews you actually found and can reference. Do NOT estimate or project.**

### 9.2 Reviews by Platform (EXACT counts only)
Create a table with the following format. Use 0 if no reviews found on a platform:
| Platform | # of Reviews | Date Range | Average Rating |
|----------|--------------|------------|----------------|
| Google Reviews | [exact count or 0] | [actual dates or N/A] | [actual avg or N/A] |
| Yelp | [exact count or 0] | [actual dates or N/A] | [actual avg or N/A] |
| Facebook | [exact count or 0] | [actual dates or N/A] | [actual avg or N/A] |
| Instagram | [exact count or 0] | [actual dates or N/A] | N/A |
| TripAdvisor | [exact count or 0] | [actual dates or N/A] | [actual avg or N/A] |
| Reddit | [exact count or 0] | [actual dates or N/A] | N/A |
| Twitter/X | [exact count or 0] | [actual dates or N/A] | N/A |
| TikTok | [exact count or 0] | [actual dates or N/A] | N/A |
| DoorDash | [exact count or 0] | [actual dates or N/A] | [actual avg or N/A] |
| UberEats | [exact count or 0] | [actual dates or N/A] | [actual avg or N/A] |
| Skip The Dishes | [exact count or 0] | [actual dates or N/A] | [actual avg or N/A] |
| YouTube | [exact count or 0] | [actual dates or N/A] | N/A |
| Local Blogs | [exact count or 0] | [actual dates or N/A] | N/A |
| **TOTAL** | **[actual total]** | | |

**Do NOT estimate review counts. If you cannot access a platform's reviews, enter 0.**

### 9.3 Rating Distribution (From Actual Reviews Only)
Count only reviews you actually found and quoted:
- **5 Stars (Perfect)**: [actual count] reviews ([actual %]%)
- **4 Stars (Excellent)**: [actual count] reviews ([actual %]%)
- **3 Stars (Mixed)**: [actual count] reviews ([actual %]%)
- **2 Stars or below (Critical)**: [actual count] reviews ([actual %]%)

**Do NOT estimate distributions based on averages. Only count reviews you have direct evidence of.**

### 9.4 Data Limitations
- List any platforms where no data was found
- Note any access restrictions encountered
- Be transparent about gaps in data coverage

## SECTION 10: SENTIMENT ANALYSIS (MANDATORY)
**THIS SECTION IS REQUIRED - Provide comprehensive sentiment analysis:**

### 10.1 Overall Sentiment Score
Calculate and provide:
- **Overall Sentiment**: [Positive/Neutral/Negative]
- **Sentiment Score**: [X.X out of 5.0]
- **Confidence Level**: [High/Medium/Low based on review volume]

### 10.2 Sentiment Breakdown
Provide percentage breakdown:
| Sentiment | Count | Percentage |
|-----------|-------|------------|
| Very Positive (Enthusiastic praise, strong recommendations) | XX | XX% |
| Positive (General satisfaction, would return) | XX | XX% |
| Neutral (Mixed feelings, balanced feedback) | XX | XX% |
| Negative (Disappointment, complaints) | XX | XX% |
| Very Negative (Strong dissatisfaction, warnings to others) | XX | XX% |

### 10.3 Sentiment by Category
Break down sentiment for each key area:
- **Product Quality Sentiment**: [Positive/Neutral/Negative] - Brief explanation
- **Service Quality Sentiment**: [Positive/Neutral/Negative] - Brief explanation
- **Value for Money Sentiment**: [Positive/Neutral/Negative] - Brief explanation
- **Atmosphere/Environment Sentiment**: [Positive/Neutral/Negative] - Brief explanation
- **Convenience/Location Sentiment**: [Positive/Neutral/Negative] - Brief explanation

### 10.4 Sentiment Trends
- **Improving**: Areas where sentiment has improved over time
- **Declining**: Areas where sentiment has worsened
- **Stable**: Areas with consistent sentiment

### 10.5 Key Sentiment Drivers
- **Top 3 Positive Drivers**: What makes customers happy
- **Top 3 Negative Drivers**: What frustrates customers
- **Net Promoter Indicator**: Based on language used (would recommend vs. would not)

## OUTPUT FORMAT:
- Be extremely detailed and thorough
- **EVERY insight MUST include a direct quote and source**
- **NEVER estimate, infer, or project data**
- If a section has no actual reviews, state "No verified reviews found"
- Organize clearly by section
- Include reviewer name, platform, date, and rating for every quote
- **ALWAYS include Section 9 with ONLY actual review counts found**
- **ALWAYS include Section 10 with sentiment analysis based ONLY on actual reviews**
- Be transparent about data gaps - it's better to show "0 reviews found" than to estimate

## FINAL REMINDER:
**QUALITY OVER QUANTITY**: A report with 10 actual verified reviews is more valuable than one with 100 estimated/inferred reviews.
**NO ESTIMATES**: Never use phrases like "estimated," "approximately," "likely," "typical," or "projected."
**CITE EVERYTHING**: Every claim needs a direct quote with source attribution.

Remember: This research will inform critical business decisions. Only verified, quotable data should be included.
"""


def extract_sentiment_data(report_content: str) -> dict:
    """Extract sentiment analysis data from the report for UI display."""
    import re

    sentiment_data = {
        'overall_sentiment': 'Positive',
        'sentiment_score': 4.0,
        'confidence': 'Medium',
        'breakdown': {
            'very_positive': {'count': 0, 'percentage': 0},
            'positive': {'count': 0, 'percentage': 0},
            'neutral': {'count': 0, 'percentage': 0},
            'negative': {'count': 0, 'percentage': 0},
            'very_negative': {'count': 0, 'percentage': 0}
        },
        'categories': {
            'product_quality': 'Positive',
            'service_quality': 'Positive',
            'value_for_money': 'Neutral',
            'atmosphere': 'Positive',
            'convenience': 'Positive'
        },
        'top_positive_drivers': [],
        'top_negative_drivers': [],
        'total_reviews': 0
    }

    try:
        # Extract overall sentiment
        overall_match = re.search(r'\*\*Overall Sentiment\*\*:\s*\[?(\w+)', report_content, re.IGNORECASE)
        if overall_match:
            sentiment_data['overall_sentiment'] = overall_match.group(1)

        # Extract sentiment score
        score_match = re.search(r'\*\*Sentiment Score\*\*:\s*\[?(\d+\.?\d*)', report_content, re.IGNORECASE)
        if score_match:
            sentiment_data['sentiment_score'] = float(score_match.group(1))

        # Extract confidence level
        confidence_match = re.search(r'\*\*Confidence Level\*\*:\s*\[?(\w+)', report_content, re.IGNORECASE)
        if confidence_match:
            sentiment_data['confidence'] = confidence_match.group(1)

        # Extract total reviews
        total_match = re.search(r'\*\*Total(?:\s+number\s+of)?\s+reviews?\s+analyzed\*\*:\s*\[?(\d+)', report_content, re.IGNORECASE)
        if total_match:
            sentiment_data['total_reviews'] = int(total_match.group(1))

        # Extract sentiment breakdown percentages
        very_pos_match = re.search(r'Very Positive.*?\|\s*(\d+)\s*\|\s*(\d+)%', report_content, re.IGNORECASE)
        if very_pos_match:
            sentiment_data['breakdown']['very_positive'] = {
                'count': int(very_pos_match.group(1)),
                'percentage': int(very_pos_match.group(2))
            }

        pos_match = re.search(r'\|\s*Positive\s*\(General.*?\|\s*(\d+)\s*\|\s*(\d+)%', report_content, re.IGNORECASE)
        if pos_match:
            sentiment_data['breakdown']['positive'] = {
                'count': int(pos_match.group(1)),
                'percentage': int(pos_match.group(2))
            }

        neutral_match = re.search(r'Neutral.*?\|\s*(\d+)\s*\|\s*(\d+)%', report_content, re.IGNORECASE)
        if neutral_match:
            sentiment_data['breakdown']['neutral'] = {
                'count': int(neutral_match.group(1)),
                'percentage': int(neutral_match.group(2))
            }

        neg_match = re.search(r'\|\s*Negative\s*\(Disappointment.*?\|\s*(\d+)\s*\|\s*(\d+)%', report_content, re.IGNORECASE)
        if neg_match:
            sentiment_data['breakdown']['negative'] = {
                'count': int(neg_match.group(1)),
                'percentage': int(neg_match.group(2))
            }

        very_neg_match = re.search(r'Very Negative.*?\|\s*(\d+)\s*\|\s*(\d+)%', report_content, re.IGNORECASE)
        if very_neg_match:
            sentiment_data['breakdown']['very_negative'] = {
                'count': int(very_neg_match.group(1)),
                'percentage': int(very_neg_match.group(2))
            }

        # Extract category sentiments
        prod_match = re.search(r'\*\*Product Quality Sentiment\*\*:\s*\[?(\w+)', report_content, re.IGNORECASE)
        if prod_match:
            sentiment_data['categories']['product_quality'] = prod_match.group(1)

        service_match = re.search(r'\*\*Service Quality Sentiment\*\*:\s*\[?(\w+)', report_content, re.IGNORECASE)
        if service_match:
            sentiment_data['categories']['service_quality'] = service_match.group(1)

        value_match = re.search(r'\*\*Value for Money Sentiment\*\*:\s*\[?(\w+)', report_content, re.IGNORECASE)
        if value_match:
            sentiment_data['categories']['value_for_money'] = value_match.group(1)

        atmosphere_match = re.search(r'\*\*Atmosphere.*?Sentiment\*\*:\s*\[?(\w+)', report_content, re.IGNORECASE)
        if atmosphere_match:
            sentiment_data['categories']['atmosphere'] = atmosphere_match.group(1)

        convenience_match = re.search(r'\*\*Convenience.*?Sentiment\*\*:\s*\[?(\w+)', report_content, re.IGNORECASE)
        if convenience_match:
            sentiment_data['categories']['convenience'] = convenience_match.group(1)

    except Exception as e:
        print(f"Error extracting sentiment data: {e}")

    return sentiment_data


def generate_word_document(bakery_location: str, report_content: str, output_path: str) -> str:
    """Generate a formatted Word document from the research report."""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('COBS Bread Bakery', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Comprehensive Review Analysis Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata section
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run('Location: ').bold = True
    meta.add_run(bakery_location)

    meta2 = doc.add_paragraph()
    meta2.add_run('Generated: ').bold = True
    meta2.add_run(datetime.now().strftime('%B %d, %Y at %I:%M %p'))

    meta3 = doc.add_paragraph()
    meta3.add_run('Research Engine: ').bold = True
    meta3.add_run('AI-Powered Deep Analysis')

    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    doc.add_paragraph()

    # Process and add the report content
    add_formatted_content(doc, report_content)

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    footer = doc.add_paragraph()
    footer.add_run('Disclaimer: ').bold = True
    footer.add_run(
        'This report is generated from publicly available reviews and social media content. '
        'All insights should be verified independently before making business decisions.'
    )

    # Save the document
    doc.save(output_path)
    return output_path


def add_formatted_content(doc: Document, content: str):
    """Add formatted content to the Word document, parsing markdown-like formatting."""
    import re
    lines = content.split('\n')

    for line in lines:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # Handle headers
        if stripped.startswith('######'):
            doc.add_heading(stripped[6:].strip(), level=6)
        elif stripped.startswith('#####'):
            doc.add_heading(stripped[5:].strip(), level=5)
        elif stripped.startswith('####'):
            doc.add_heading(stripped[4:].strip(), level=4)
        elif stripped.startswith('###'):
            doc.add_heading(stripped[3:].strip(), level=3)
        elif stripped.startswith('##'):
            doc.add_heading(stripped[2:].strip(), level=2)
        elif stripped.startswith('#'):
            doc.add_heading(stripped[1:].strip(), level=1)

        # Handle bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')

        # Handle numbered lists
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
            doc.add_paragraph(stripped[2:].strip(), style='List Number')

        # Handle bold text markers
        elif stripped.startswith('**') and stripped.endswith('**'):
            para = doc.add_paragraph()
            para.add_run(stripped[2:-2]).bold = True

        # Regular paragraph
        else:
            para = doc.add_paragraph()
            # Handle inline bold
            pattern = r'\*\*([^*]+)\*\*'
            parts = re.split(pattern, stripped)

            for i, part in enumerate(parts):
                if i % 2 == 0:
                    para.add_run(part)
                else:
                    para.add_run(part).bold = True


def run_research(task_id: str, location: str):
    """Background task to run the research."""
    try:
        update_task(task_id, {'status': 'running', 'stage': 'prefetch_reviews'})

        # =================================================================
        # STAGE 1: Prefetch Google Reviews (5 most recent)
        # =================================================================
        print(f"[{task_id}] Stage 1: Fetching Google Reviews...")
        google_reviews = fetch_google_reviews(location)

        if google_reviews.get('success'):
            print(f"[{task_id}] Got {len(google_reviews.get('reviews', []))} Google reviews")
            update_task(task_id, {
                'stage': 'prefetch_search',
                'google_reviews_count': len(google_reviews.get('reviews', [])),
                'google_rating': google_reviews.get('rating')
            })
        else:
            print(f"[{task_id}] Google Reviews fetch failed: {google_reviews.get('error')}")
            update_task(task_id, {'stage': 'prefetch_search', 'google_reviews_error': google_reviews.get('error')})

        # =================================================================
        # STAGE 2: Prefetch Search Grounding Insights (multi-platform)
        # =================================================================
        print(f"[{task_id}] Stage 2: Fetching Search Grounding insights...")
        search_insights = fetch_search_grounding_insights(location)

        if search_insights.get('success'):
            print(f"[{task_id}] Got search insights from {len(search_insights.get('sources', []))} sources")
            update_task(task_id, {
                'stage': 'deep_research',
                'search_sources_count': len(search_insights.get('sources', []))
            })
        else:
            print(f"[{task_id}] Search Grounding failed: {search_insights.get('error')}")
            update_task(task_id, {'stage': 'deep_research', 'search_error': search_insights.get('error')})

        # =================================================================
        # STAGE 3: Run Deep Research with prefetched data
        # =================================================================
        print(f"[{task_id}] Stage 3: Starting Deep Research...")

        # Initialize the Google client
        client = genai.Client()

        # Create the research interaction with prefetched data
        prompt = build_research_prompt(location, google_reviews, search_insights)
        interaction = client.interactions.create(
            input=prompt,
            agent=AGENT_ID,
            background=True
        )

        update_task(task_id, {'interaction_id': interaction.id})

        # Poll for results
        max_poll_time = 3600  # 60 minutes
        poll_interval = 15
        start_time = time.time()

        while True:
            elapsed = time.time() - start_time

            if elapsed > max_poll_time:
                update_task(task_id, {
                    'status': 'failed',
                    'error': 'Research exceeded maximum time limit'
                })
                return

            interaction = client.interactions.get(interaction.id)

            if interaction.status == "completed":
                if interaction.outputs:
                    report = interaction.outputs[-1].text

                    # Generate Word document
                    OUTPUTS_DIR.mkdir(exist_ok=True)
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    safe_location = "".join(c if c.isalnum() or c in ' -_' else '_' for c in location)[:50]
                    doc_path = OUTPUTS_DIR / f"COBS_Research_{safe_location}_{timestamp}.docx"

                    generate_word_document(location, report, str(doc_path))

                    # Extract sentiment data for UI display
                    sentiment_data = extract_sentiment_data(report)

                    update_task(task_id, {
                        'status': 'completed',
                        'report': report,
                        'report_length': len(report),
                        'document_path': str(doc_path),
                        'sentiment': sentiment_data
                    })
                else:
                    update_task(task_id, {
                        'status': 'failed',
                        'error': 'Research completed but no output received'
                    })
                return

            elif interaction.status == "failed":
                update_task(task_id, {
                    'status': 'failed',
                    'error': getattr(interaction, 'error', 'Unknown error')
                })
                return

            time.sleep(poll_interval)

    except Exception as e:
        update_task(task_id, {
            'status': 'failed',
            'error': str(e)
        })


# Routes
@app.route('/')
def index():
    """Serve the main page."""
    return render_template('index.html')


@app.route('/favicon.ico')
def favicon():
    """Serve favicon - COBS brand color bread icon."""
    # SVG favicon with COBS brand color
    svg = '''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 32 32">
        <circle cx="16" cy="16" r="15" fill="#862633"/>
        <ellipse cx="16" cy="16" rx="10" ry="7" fill="#D4A574"/>
        <ellipse cx="16" cy="14" rx="8" ry="5" fill="#E8CDB0"/>
        <path d="M10 16 Q16 12 22 16" stroke="#862633" stroke-width="1" fill="none"/>
        <path d="M12 15 Q16 11 20 15" stroke="#862633" stroke-width="0.5" fill="none"/>
    </svg>'''
    return Response(svg, mimetype='image/svg+xml')


@app.route('/api/research', methods=['POST'])
def start_research():
    """Start a new research task."""
    data = request.get_json()
    location = data.get('location', '').strip()

    if not location:
        return jsonify({'error': 'Location is required'}), 400

    # Check for API key
    if not os.environ.get('GOOGLE_API_KEY'):
        return jsonify({'error': 'Google API key not configured. Please add GOOGLE_API_KEY environment variable.'}), 500

    # Create task
    task_id = str(uuid.uuid4())
    task_data = {
        'id': task_id,
        'location': location,
        'status': 'pending',
        'created_at': datetime.now().isoformat(),
        'report': None,
        'document_path': None,
        'error': None
    }
    create_task(task_id, task_data)

    # Start background task
    thread = threading.Thread(target=run_research, args=(task_id, location))
    thread.daemon = True
    thread.start()

    return jsonify({
        'task_id': task_id,
        'status': 'pending',
        'message': 'Research started'
    })


@app.route('/api/research/<task_id>')
def get_research_status(task_id):
    """Get the status of a research task."""
    task = get_task(task_id)

    if not task:
        return jsonify({'error': 'Task not found. The task may have expired or the server was restarted.'}), 404

    response = {
        'task_id': task_id,
        'status': task['status'],
        'location': task['location']
    }

    if task['status'] == 'completed':
        response['report_length'] = task.get('report_length', 0)
        response['document_path'] = task.get('document_path')
        response['sentiment'] = task.get('sentiment', {})

    if task['status'] == 'failed':
        response['error'] = task.get('error', 'Unknown error')

    return jsonify(response)


@app.route('/api/download/<task_id>')
def download_document(task_id):
    """Download the generated Word document."""
    task = get_task(task_id)

    if not task:
        return jsonify({'error': 'Task not found'}), 404

    if task['status'] != 'completed':
        return jsonify({'error': 'Research not completed'}), 400

    doc_path = task.get('document_path')
    if not doc_path or not Path(doc_path).exists():
        return jsonify({'error': 'Document not found'}), 404

    return send_file(
        doc_path,
        as_attachment=True,
        download_name=Path(doc_path).name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/health')
def health():
    """Health check endpoint."""
    return jsonify({'status': 'healthy'})


if __name__ == '__main__':
    # Create outputs directory
    OUTPUTS_DIR.mkdir(exist_ok=True)

    # Run the app
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_DEBUG', 'false').lower() == 'true'
    app.run(host='0.0.0.0', port=port, debug=debug)
